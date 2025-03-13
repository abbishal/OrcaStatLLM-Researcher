from md2pdf.core import md2pdf
from pathlib import Path
from modules.utils.async_buffer import AsyncBuffer
import os
import base64
import tempfile
import shutil
import re
import subprocess
from docx import Document
ACADEMIC_PAPER_CSS = """
body {
    font-family: 'Times New Roman', Times, serif;
    font-size: 12pt;
    line-height: 1.5;
    margin: 1in;
    color: #333;
}
h1, h2, h3, h4, h5, h6 {
    font-family: 'Arial', sans-serif;
    color: #000;
    margin-top: 0.8em;
    margin-bottom: 0.3em;
}
h1 {
    font-size: 18pt;
    text-align: center;
    margin-top: 1.5em;
    margin-bottom: 0.8em;
}
h2 {
    font-size: 14pt;
    border-bottom: 1px solid #ddd;
    padding-bottom: 0.1em;
    margin-top: 1em;
}
h3 {
    font-size: 12pt;
    font-weight: bold;
}
p {
    margin-bottom: 0.5em;
    text-align: justify;
}
img {
    max-width: 90%;
    height: auto;
    display: block;
    margin: 0.5em auto;
    page-break-inside: avoid;
}
figure {
    margin: 1em 0;
    page-break-inside: avoid;
}
table {
    width: 95%;
    border-collapse: collapse;
    margin: 1em auto;
    page-break-inside: avoid;
    font-size: 10pt;
    border: 1px solid #ddd;
}
th, td {
    border: 1px solid #ddd;
    padding: 0.4em 0.6em;
    text-align: left;
    vertical-align: top;
}
th {
    background-color: #f2f2f2;
    font-weight: bold;
}
blockquote {
    border-left: 4px solid #ddd;
    padding-left: 1em;
    margin-left: 0;
    color: #555;
}
code {
    font-family: 'Courier New', monospace;
    background-color: #f5f5f5;
    padding: 0.2em 0.4em;
    border-radius: 3px;
    font-size: 0.9em;
}
sup {
    font-size: 75%;
    vertical-align: super;
}
a {
    color: #0366d6;
    text-decoration: none;
}
.caption {
    text-align: center;
    font-style: italic;
    margin-top: 0.3em;
    margin-bottom: 0.5em;
    color: #666;
    font-size: 10pt;
}
.abstract {
    font-style: italic;
    margin: 1em 2em;
}
.footnotes {
    font-size: 9pt;
    border-top: 1px solid #ddd;
    margin-top: 1em;
    padding-top: 0.5em;
}
.pagebreak {
    page-break-before: always;
    height: 0;
    display: block;
}
    page-break-after: always;
}
.logo-container {
    text-align: center;
    margin-bottom: 1em;
    margin-top: 1em;
}
/* Specific table styles to ensure proper rendering in PDF */
.markdown-table {
    width: 100%;
    border: 1px solid #ddd;
    border-spacing: 0;
    border-collapse: collapse;
}
.markdown-table td, .markdown-table th {
    border: 1px solid #ddd;
    padding: 6px 13px;
}
.markdown-table th {
    background-color: #f2f2f2;
}
"""

class PDFConverter:
    async def convert_to_pdf(self, markdown_file: str, buffer: AsyncBuffer) -> str:
        try:
            buffer.add_log(f"Converting markdown to PDF: {markdown_file}", high_level=True)
            css_file = self._prepare_css_file()
            buffer.add_log(f"Created academic CSS styling", high_level=True)
            temp_md_file = self._preprocess_markdown(markdown_file)
            buffer.add_log(f"Preprocessed markdown for improved PDF formatting", high_level=True)
            
            pdf_file = markdown_file.replace('.md', '.pdf')
            
            try:
                md2pdf(
                    pdf_file,
                    md_file_path=temp_md_file,
                    css_file_path=css_file,
                    base_url=str(Path(markdown_file).parent)
                )
                
                buffer.add_log(f"PDF file generated: {pdf_file}", high_level=True)
                if os.path.exists(temp_md_file):
                    os.remove(temp_md_file)
                if os.path.exists(css_file):
                    os.remove(css_file)
                    
                return pdf_file
            except Exception as e:
                buffer.add_log(f"Error generating PDF with md2pdf: {str(e)}", high_level=True)
                buffer.add_log("Trying alternative method...", high_level=True)
                alternative_result = self._convert_with_pandoc(temp_md_file, pdf_file, buffer)
                if alternative_result:
                    return pdf_file
                else:
                    buffer.add_log("All conversion methods failed", high_level=True)
                    return ""
                
        except Exception as e:
            buffer.add_log(f"Error converting to PDF: {str(e)}", high_level=True)
            return ""
    
    async def convert_to_docx(self, markdown_file: str, buffer: AsyncBuffer) -> str:

        try:
            buffer.add_log(f"Converting markdown to DOCX: {markdown_file}", high_level=True)
            temp_md_file = self._preprocess_markdown(markdown_file)
            buffer.add_log(f"Preprocessed markdown for improved DOCX formatting", high_level=True)
            
            docx_file = markdown_file.replace('.md', '.docx')
            try:
                cmd = ['pandoc', temp_md_file, '-o', docx_file, '--reference-doc=reference.docx']
                reference_file = Path(markdown_file).parent / "reference.docx"
                if not reference_file.exists():
                    cmd = ['pandoc', temp_md_file, '-o', docx_file]
                
                process = subprocess.run(cmd, capture_output=True, text=True)
                
                if process.returncode == 0:
                    buffer.add_log(f"DOCX file generated: {docx_file}", high_level=True)
                    if os.path.exists(temp_md_file):
                        os.remove(temp_md_file)
                    
                    return docx_file
                else:
                    buffer.add_log(f"Error generating DOCX with pandoc: {process.stderr}", high_level=True)
                    raise Exception(f"Pandoc error: {process.stderr}")
                    
            except Exception as e:
                buffer.add_log(f"Error with pandoc conversion: {str(e)}", high_level=True)
                buffer.add_log("Falling back to python-docx method...", high_level=True)
                return self._convert_with_python_docx(markdown_file, buffer)
                
        except Exception as e:
            buffer.add_log(f"Error converting to DOCX: {str(e)}", high_level=True)
            return ""
    
    def _convert_with_python_docx(self, markdown_file: str, buffer: AsyncBuffer) -> str:

        try:
            with open(markdown_file, 'r', encoding='utf-8') as f:
                content = f.read()
            doc = Document()
            lines = content.split('\n')
            current_paragraph = None
            
            for line in lines:
                if line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                elif line.strip() == '':
                    current_paragraph = None
                else:
                    if current_paragraph is None:
                        current_paragraph = doc.add_paragraph()
                    current_paragraph.add_run(line)
            docx_file = markdown_file.replace('.md', '.docx')
            doc.save(docx_file)
            
            buffer.add_log(f"DOCX file generated with python-docx: {docx_file}", high_level=True)
            return docx_file
            
        except Exception as e:
            buffer.add_log(f"Error in python-docx conversion: {str(e)}", high_level=True)
            return ""

    def _prepare_css_file(self) -> str:

        css_file = tempfile.NamedTemporaryFile(delete=False, suffix='.css')
        with open(css_file.name, 'w') as f:
            f.write(ACADEMIC_PAPER_CSS)
        return css_file.name

    def _preprocess_markdown(self, markdown_file: str) -> str:

        with open(markdown_file, 'r', encoding='utf-8') as f:
            content = f.read()
        content = self._add_logo_to_document(content)
        content = self._fix_image_paths(content, Path(markdown_file).parent)
        content = self._fix_tables(content)
        content = self._fix_subtopic_formatting(content)
        content = self._fix_toc_and_anchors(content)
        content = self._add_page_breaks(content)
        content = self._remove_duplicate_images(content)
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.md')
        with open(temp_file.name, 'w', encoding='utf-8') as f:
            f.write(content)
        
        return temp_file.name
    
    def _fix_subtopic_formatting(self, content: str) -> str:

        pattern = r'## ([^\n]+)\s+\n+([^\n]+)\1'
        content = re.sub(pattern, r'## \1', content)
        pattern2 = r'## ([^\n]+)\s+\n+### \1'
        content = re.sub(pattern2, r'## \1', content)
        
        return content

    def _remove_duplicate_images(self, content: str) -> str:

        duplicate_pattern = r'<figure>[\s\S]*?<img src="([^"]+)"[\s\S]*?<\/figure>\s*\n\s*!\[[^\]]*\]\(\1\)'
        return re.sub(duplicate_pattern, r'<figure>\n<img src="\1" alt="Figure" style="max-width:90%; margin:0 auto; display:block;">\n<figcaption style="text-align:center; font-style:italic;">Figure</figcaption>\n</figure>', content)

    def _add_logo_to_document(self, content: str) -> str:

        logo_svg = """<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 400 400" width="150" height="150">
  <!-- Background neural network pattern -->
  <defs>
    <pattern id="neuralPattern" x="0" y="0" width="100" height="100" patternUnits="userSpaceOnUse">
      <circle cx="10" cy="10" r="1.5" fill="#3498db" opacity="0.3"/>
      <circle cx="30" cy="20" r="1.5" fill="#9b59b6" opacity="0.3"/>
      <circle cx="50" cy="10" r="1.5" fill="#3498db" opacity="0.3"/>
      <circle cx="70" cy="30" r="1.5" fill="#9b59b6" opacity="0.3"/>
      <circle cx="90" cy="20" r="1.5" fill="#3498db" opacity="0.3"/>
      <circle cx="20" cy="40" r="1.5" fill="#9b59b6" opacity="0.3"/>
      <circle cx="40" cy="50" r="1.5" fill="#3498db" opacity="0.3"/>
      <circle cx="60" cy="40" r="1.5" fill="#9b59b6" opacity="0.3"/>
      <circle cx="80" cy="60" r="1.5" fill="#3498db" opacity="0.3"/>
      <circle cx="10" cy="70" r="1.5" fill="#9b59b6" opacity="0.3"/>
      <circle cx="30" cy="80" r="1.5" fill="#3498db" opacity="0.3"/>
      <circle cx="50" cy="70" r="1.5" fill="#9b59b6" opacity="0.3"/>
      <circle cx="70" cy="90" r="1.5" fill="#3498db" opacity="0.3"/>
      <circle cx="90" cy="80" r="1.5" fill="#9b59b6" opacity="0.3"/>
      <line x1="10" y1="10" x2="30" y2="20" stroke="#3498db" stroke-width="0.5" opacity="0.2"/>
      <line x1="30" y1="20" x2="50" y2="10" stroke="#9b59b6" stroke-width="0.5" opacity="0.2"/>
      <line x1="50" y1="10" x2="70" y2="30" stroke="#3498db" stroke-width="0.5" opacity="0.2"/>
      <line x1="70" y1="30" x2="90" y2="20" stroke="#9b59b6" stroke-width="0.5" opacity="0.2"/>
      <line x1="20" y1="40" x2="40" y2="50" stroke="#3498db" stroke-width="0.5" opacity="0.2"/>
      <line x1="40" y1="50" x2="60" y2="40" stroke="#9b59b6" stroke-width="0.5" opacity="0.2"/>
      <line x1="60" y1="40" x2="80" y2="60" stroke="#3498db" stroke-width="0.5" opacity="0.2"/>
      <line x1="10" y1="70" x2="30" y2="80" stroke="#9b59b6" stroke-width="0.5" opacity="0.2"/>
      <line x1="30" y1="80" x2="50" y2="70" stroke="#3498db" stroke-width="0.5" opacity="0.2"/>
      <line x1="50" y1="70" x2="70" y2="90" stroke="#9b59b6" stroke-width="0.5" opacity="0.2"/>
      <line x1="70" y1="90" x2="90" y2="80" stroke="#3498db" stroke-width="0.5" opacity="0.2"/>
    </pattern>
    
    <!-- Glitch filters -->
    <filter id="glitchFilter">
      <feTurbulence type="fractalNoise" baseFrequency="0.05" numOctaves="2" result="noise"/>
      <feDisplacementMap in="SourceGraphic" in2="noise" scale="3" xChannelSelector="R" yChannelSelector="G"/>
    </filter>
    
    <!-- Gradients -->
    <linearGradient id="nameGradient" x1="0%" y1="0%" x2="100%" y2="0%">
      <stop offset="0%" stop-color="#3498db"/>
      <stop offset="100%" stop-color="#9b59b6"/>
    </linearGradient>
    <linearGradient id="labsGradient" x1="0%" y1="0%" x2="100%" y2="0%">
      <stop offset="0%" stop-color="#9b59b6"/>
      <stop offset="100%" stop-color="#3498db"/>
    </linearGradient>
    
    <!-- Hazy effect -->
    <filter id="hazyEffect">
      <feGaussianBlur stdDeviation="0.5" result="blur"/>
      <feBlend in="SourceGraphic" in2="blur" mode="normal"/>
    </filter>
  </defs>
  
  <!-- Background with neural network pattern -->
  <rect x="0" y="0" width="400" height="400" fill="url(#neuralPattern)" opacity="0.5"/>
  
  <!-- Modern square frame with glitchy effect -->
  <rect x="40" y="40" width="320" height="320" fill="none" stroke="#333" stroke-width="6" rx="15" ry="15" filter="url(#glitchFilter)"/>
  
  <!-- Neural network nodes and connections -->
  <g opacity="0.7">
    <!-- Layer 1 nodes -->
    <circle cx="100" cy="280" r="5" fill="#3498db"/>
    <circle cx="140" cy="280" r="5" fill="#3498db"/>
    <circle cx="180" cy="280" r="5" fill="#3498db"/>
    <circle cx="220" cy="280" r="5" fill="#3498db"/>
    <circle cx="260" cy="280" r="5" fill="#3498db"/>
    <circle cx="300" cy="280" r="5" fill="#3498db"/>
    
    <!-- Layer 2 nodes -->
    <circle cx="130" cy="250" r="5" fill="#9b59b6"/>
    <circle cx="200" cy="250" r="5" fill="#9b59b6"/>
    <circle cx="270" cy="250" r="5" fill="#9b59b6"/>
    
    <!-- Layer 3 nodes -->
    <circle cx="200" cy="220" r="5" fill="#3498db"/>
    
    <!-- Connections -->
    <line x1="100" y1="280" x2="130" y2="250" stroke="#3498db" stroke-width="1"/>
    <line x1="140" y1="280" x2="130" y2="250" stroke="#3498db" stroke-width="1"/>
    <line x1="140" y1="280" x2="200" y2="250" stroke="#3498db" stroke-width="1"/>
    <line x1="180" y1="280" x2="200" y2="250" stroke="#3498db" stroke-width="1"/>
    <line x1="220" y1="280" x2="200" y2="250" stroke="#3498db" stroke-width="1"/>
    <line x1="220" y1="280" x2="270" y2="250" stroke="#3498db" stroke-width="1"/>
    <line x1="260" y1="280" x2="270" y2="250" stroke="#3498db" stroke-width="1"/>
    <line x1="300" y1="280" x2="270" y2="250" stroke="#3498db" stroke-width="1"/>
    <line x1="130" y1="250" x2="200" y2="220" stroke="#9b59b6" stroke-width="1"/>
    <line x1="200" y1="250" x2="200" y2="220" stroke="#9b59b6" stroke-width="1"/>
    <line x1="270" y1="250" x2="200" y2="220" stroke="#9b59b6" stroke-width="1"/>
  </g>
  
  <!-- Glitchy text effects -->
  <g filter="url(#hazyEffect)">
    <!-- Glitchy copy of AlgoNet text -->
    <text x="198" y="178" font-family="'Courier New', monospace" font-weight="700" font-size="42" text-anchor="middle" fill="rgba(157, 89, 182, 0.3)">AlgoNet</text>
    <!-- Main AlgoNet text -->
    <text x="200" y="180" font-family="'Courier New', monospace" font-weight="700" font-size="42" text-anchor="middle" fill="url(#nameGradient)">AlgoNet</text>
    
    <!-- Glitchy copy of Labs text -->
    <text x="238" y="228" font-family="'Courier New', monospace" font-weight="700" font-size="42" text-anchor="middle" fill="rgba(52, 152, 219, 0.3)">Labs</text>
    <!-- Main Labs text -->
    <text x="240" y="230" font-family="'Courier New', monospace" font-weight="700" font-size="42" text-anchor="middle" fill="url(#labsGradient)">Labs</text>
  </g>
  
  <!-- Small tagline with glitch effect -->
  <text x="200" y="310" font-family="'Courier New', monospace" font-size="14" text-anchor="middle" fill="#555" filter="url(#glitchFilter)">RAG · LLM · AI RESEARCH</text>
  
  <!-- Additional glitchy elements -->
  <rect x="60" y="60" width="10" height="5" fill="#3498db" opacity="0.5"/>
  <rect x="330" y="70" width="10" height="5" fill="#9b59b6" opacity="0.5"/>
  <rect x="320" y="330" width="10" height="5" fill="#3498db" opacity="0.5"/>
  <rect x="70" y="330" width="10" height="5" fill="#9b59b6" opacity="0.5"/>
</svg>"""
        logo_html = f'<div class="logo-container">{logo_svg}</div>\n\n'
        title_match = re.search(r'# .*', content)
        if title_match:
            title_pos = title_match.start()
            content = content[:title_pos] + logo_html + content[title_pos:]
        else:
            content = logo_html + content
        
        return content

    def _fix_image_paths(self, content: str, base_dir: Path) -> str:

        img_pattern = r'!\[(.*?)\]\((.*?)\)'
        
        def process_image_path(match):
            alt_text = match.group(1)
            path = match.group(2)
            if path.startswith('data:image'):
                try:
                    img_type = path.split(';')[0].split('/')[1]
                    img_data = path.split(',')[1]
                    filename = f"image_{hash(img_data) % 10000}.{img_type}"
                    img_path = base_dir / filename
                    
                    with open(img_path, 'wb') as f:
                        f.write(base64.b64decode(img_data))
                    
                    return f"![{alt_text}]({filename})"
                except Exception:
                    return match.group(0)  # Return original if extraction fails
            if path.startswith('figures/'):
                direct_path = base_dir / path
                
                if os.path.exists(direct_path):
                    return match.group(0)
                else:
                    alt_path = base_dir / path.replace('figures/', '')
                    if os.path.exists(alt_path):
                        figures_dir = base_dir / "figures"
                        figures_dir.mkdir(exist_ok=True)
                        
                        dest_path = figures_dir / os.path.basename(alt_path)
                        try:
                            shutil.copy2(alt_path, dest_path)
                            return f"![{alt_text}]({path})"
                        except Exception:
                            return f"![{alt_text}]({alt_path.name})"
                    else:
                        print(f"WARNING: Image not found: {path}")
                        return match.group(0)
            else:
                img_path = base_dir / path
                if os.path.exists(img_path):
                    return match.group(0)
                else:
                    figures_path = base_dir / "figures" / path
                    if os.path.exists(figures_path):
                        return f"![{alt_text}](figures/{path})"
                    else:
                        print(f"WARNING: Image not found: {path}")
                        return match.group(0)
        return re.sub(img_pattern, process_image_path, content)
    
    def _fix_tables(self, content: str) -> str:

        table_pattern = r'(\|[^\n]+\|\n\|[-:| ]+\|\n(?:\|[^\n]+\|\n)+)'
        
        def process_table(match):
            table = match.group(0)
            lines = table.strip().split('\n')
            if len(lines) < 3:  # Need at least header, separator, and one data row
                return table
            header_cells = [cell.strip() for cell in lines[0].split('|')]
            header_cells = [cell for cell in header_cells if cell]  # Remove empty cells
            separator_cells = [cell.strip() for cell in lines[1].split('|')]
            separator_cells = [cell for cell in separator_cells if cell]
            
            alignments = []
            for sep in separator_cells:
                if sep.startswith(':') and sep.endswith(':'):
                    alignments.append('center')
                elif sep.endswith(':'):
                    alignments.append('right')
                else:
                    alignments.append('left')
            while len(alignments) < len(header_cells):
                alignments.append('left')
            html_table = '<div class="table-wrapper">\n<table class="markdown-table">\n<thead>\n<tr>\n'
            for i, cell in enumerate(header_cells):
                align = alignments[i] if i < len(alignments) else 'left'
                html_table += f'<th style="text-align: {align};">{cell}</th>\n'
            
            html_table += '</tr>\n</thead>\n<tbody>\n'
            for i in range(2, len(lines)):
                row_cells = [cell.strip() for cell in lines[i].split('|')]
                row_cells = [cell for cell in row_cells if cell]  # Remove empty cells
                
                if not row_cells:  # Skip empty rows
                    continue
                    
                html_table += '<tr>\n'
                for j, cell in enumerate(row_cells):
                    align = alignments[j] if j < len(alignments) else 'left'
                    html_table += f'<td style="text-align: {align};">{cell}</td>\n'
                html_table += '</tr>\n'
            
            html_table += '</tbody>\n</table>\n</div>\n'
            
            return html_table
        content = re.sub(table_pattern, process_table, content)
        return content
    
    def _add_page_breaks(self, content: str) -> str:

        content = re.sub(
            r'(# .+?\n\n\*.+?\*)\n\n',
            r'\1\n\n<div class="pagebreak"></div>\n\n',
            content
        )
        content = re.sub(
            r'(## Table of Contents\n(?:.*\n)+?)\n(##\s)',
            r'\1\n<div class="pagebreak"></div>\n\n\2',
            content
        )
        content = content.replace('## Table of Contents', '## Table of Contents {#table-of-contents}')
        content = re.sub(
            r'(\n## References)',
            r'\n<div class="pagebreak"></div>\1',
            content
        )
        
        return content
    
    def _enhance_figure_captions(self, content: str) -> str:

        img_caption_pattern = r'(!\[.*?\]\(.*?\))\s*\n\s*(\*.*?\*|<div[^>]*?>.*?</div>)'
        
        def enhance_caption(match):
            img_ref = match.group(1)
            caption = match.group(2)
            if '*Figure:' not in caption and '<div' not in caption:
                if caption.startswith('*') and caption.endswith('*'):
                    enhanced_caption = caption.replace('*', '*Figure: ', 1)
                else:
                    enhanced_caption = f"*Figure: {caption.strip('* ')}*"
                html_caption = f'<div class="caption">{enhanced_caption[1:-1]}</div>'
                
                return f"{img_ref}\n\n{html_caption}\n"
            else:
                return f"{img_ref}\n\n{caption}\n"
        return re.sub(img_caption_pattern, enhance_caption, content)

    def _convert_with_pandoc(self, md_file: str, pdf_file: str, buffer: AsyncBuffer) -> str:

        try:
            import subprocess
            
            buffer.add_log("Attempting PDF conversion with pandoc", high_level=True)
            try:
                subprocess.run(['pandoc', '--version'], check=True, capture_output=True)
            except (subprocess.SubprocessError, FileNotFoundError):
                buffer.add_log("Pandoc not found, cannot use alternative method", high_level=True)
                return None
            latex_template = self._create_latex_template()
            cmd = [
                'pandoc',
                md_file,
                '-o', pdf_file,
                '--pdf-engine=xelatex',
                '--template=' + latex_template,
                '--variable', 'geometry:margin=0.75in',
                '--variable', 'fontsize=12pt',
                '--variable', 'colorlinks=true',
                '--toc'
            ]
            
            result = subprocess.run(cmd, check=True, capture_output=True)
            if os.path.exists(latex_template):
                os.remove(latex_template)
            
            if os.path.exists(pdf_file):
                buffer.add_log("Successfully converted to PDF using pandoc", high_level=True)
                return pdf_file
            else:
                buffer.add_log(f"Pandoc ran but did not produce PDF output", high_level=True)
                return None
                
        except Exception as e:
            buffer.add_log(f"Error using pandoc: {str(e)}", high_level=True)
            return None

    def _create_latex_template(self) -> str:

        template = r"""
\documentclass[12pt]{article}
\usepackage{geometry}
\geometry{margin=0.75in}
\usepackage{fontspec}
\usepackage{hyperref}
\usepackage{graphicx}
\usepackage{booktabs}
\usepackage[table]{xcolor}
\usepackage{caption}
\usepackage{float}
\usepackage{fancyhdr}

% Improve table formatting
\usepackage{array}
\newcolumntype{L}[1]{>{\raggedright\arraybackslash}p{#1}}
\newcolumntype{C}[1]{>{\centering\arraybackslash}p{#1}}
\newcolumntype{R}[1]{>{\raggedleft\arraybackslash}p{#1}}

% Improve image handling
\usepackage{adjustbox}

% Set default figure placement to H (strictly here)
\makeatletter
\def\fps@figure{H}
\makeatother

% Configure hyperlinks
\hypersetup{
    colorlinks=true,
    linkcolor=blue,
    filecolor=magenta,
    urlcolor=blue,
}

\begin{document}
$body$
\end{document}
"""
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.tex')
        with open(temp_file.name, 'w', encoding='utf-8') as f:
            f.write(template)
        return temp_file.name

    def _fix_toc_and_anchors(self, content: str) -> str:

        headings = re.findall(r'^#{1,6}\s+(.*?)$', content, re.MULTILINE)
        anchor_dict = {}
        for heading in headings:
            clean_heading = heading.strip()
            anchor = clean_heading.lower().replace(' ', '-')
            anchor = re.sub(r'[^\w\-]', '', anchor)
            anchor_dict[clean_heading] = anchor
            heading_with_anchor = f'<a id="{anchor}"></a>\n## {clean_heading}'
            content = content.replace(f'## {clean_heading}', heading_with_anchor)
        toc_section = re.search(r'## Table of Contents\s+([\s\S]*?)(?=^##\s+(?!Table of Contents)|\Z)', 
                               content, re.MULTILINE)
        
        if toc_section:
            toc_content = toc_section.group(1)
            toc_entries = re.findall(r'\d+\.\s+\[(.*?)\]\(#(.*?)\)', toc_content)
            
            for entry_text, entry_anchor in toc_entries:
                if entry_text in anchor_dict:
                    correct_anchor = anchor_dict[entry_text]
                    old_link = f'[{entry_text}](#{entry_anchor})'
                    new_link = f'[{entry_text}](#{correct_anchor})'
                    content = content.replace(old_link, new_link)
        
        return content

